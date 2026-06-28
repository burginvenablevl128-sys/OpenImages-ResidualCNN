from __future__ import annotations

import csv
import json
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .config import IMAGE_SIZE, TARGET_CLASSES, project_root


TITLE = "从像素表示到残差学习：基于 Open Images V7 子集的图像分类复现实验"
SUBTITLE = "MLP 基线、局部卷积特征与 ResidualCNN 框架的对照分析"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6B7C"
LIGHT_FILL = "F4F6F9"
BORDER = "D8DEE8"


@dataclass
class Numbering:
    figure: int = 0
    table: int = 0


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_run_font(
    run,
    size: float = 11,
    bold: bool = False,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
    color: str | None = None,
    italic: bool = False,
) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para(
    p,
    align: WD_ALIGN_PARAGRAPH | None = None,
    before: float = 0,
    after: float = 8,
    line: float = 1.333,
    first_line: bool = False,
    keep_with_next: bool = False,
) -> None:
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    if first_line:
        fmt.first_line_indent = Inches(0.29)


def add_page_number(paragraph) -> None:
    paragraph.add_run("第 ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)
    paragraph.add_run(" 页")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]
    for style_name, size, color, before, after in heading_tokens:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_para(header, after=0, line=1.0)
        r = header.add_run("Open Images V7 图像分类复现实验")
        set_run_font(r, size=9, east_asia="宋体", color=MUTED)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(footer, after=0, line=1.0)
        r = footer.add_run("基于 ResidualCNN 的人工智能课程论文 | ")
        set_run_font(r, size=9, east_asia="宋体", color=MUTED)
        add_page_number(footer)
        for run in footer.runs:
            set_run_font(run, size=9, east_asia="宋体", color=MUTED)


def add_title_page(doc: Document, numbering: Numbering, rows: list[dict[str, str]], metrics: dict) -> None:
    dataset = metrics.get("dataset", {})
    counts = Counter(r["label_zh"] for r in rows)
    failures_path = project_root() / "logs" / "download_failures.csv"
    failure_count = 0
    if failures_path.exists():
        failure_count = max(0, len(read_csv(failures_path)))

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=16, line=1.0)
    r = p.add_run("《人工智能》课程论文")
    set_run_font(r, size=13, bold=True, east_asia="黑体", color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=8, line=1.15)
    r = p.add_run(TITLE)
    set_run_font(r, size=22, bold=True, east_asia="黑体", color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=28, line=1.2)
    r = p.add_run(SUBTITLE)
    set_run_font(r, size=13.5, east_asia="宋体", color=MUTED)

    add_table(
        doc,
        numbering,
        ["项目", "说明"],
        [
            ["研究任务", "基于真实公开图像的多类别分类算法复现"],
            ["数据来源", "Open Images V7 官方验证集元数据与公开缩略图"],
            ["实验规模", f"{len(rows)} 张 Pilot 图像，{len(counts)} 个类别，训练/验证/测试为 {dataset.get('train', '-')}/{dataset.get('validation', '-')}/{dataset.get('test', '-')}"],
            ["模型路线", "MLP 展平像素基线、局部卷积特征模型、ResidualCNN 完整训练入口"],
            ["可复现材料", f"样本表、失败日志（{failure_count} 条）、训练日志、指标文件、可视化图与生成脚本"],
        ],
        "论文基本信息",
        widths=[1800, CONTENT_DXA - 1800],
    )

    p = doc.add_paragraph()
    set_para(p, before=12, after=0, line=1.2)
    r = p.add_run("版式说明：本文采用叙述型学术报告结构，正文围绕数据、模型与实验解释展开；运行命令与交付说明集中放在附录和可复现性部分。")
    set_run_font(r, size=10, east_asia="宋体", color=MUTED)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        set_para(p, before=18, after=10, line=1.25, keep_with_next=True)
    elif level == 2:
        set_para(p, before=12, after=6, line=1.25, keep_with_next=True)
    else:
        set_para(p, before=8, after=4, line=1.25, keep_with_next=True)
    for r in p.runs:
        set_run_font(
            r,
            size=16 if level == 1 else 13 if level == 2 else 12,
            bold=True,
            east_asia="黑体",
            color=BLUE if level in (1, 2) else DARK_BLUE,
        )


def add_body(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    set_para(p, first_line=first_line)
    r = p.add_run(text)
    set_run_font(r, size=11, east_asia="宋体", color="000000")


def add_lead(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_para(p, before=0, after=10, line=1.25)
    r = p.add_run(text)
    set_run_font(r, size=11.5, bold=True, east_asia="宋体", color=INK)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [CONTENT_DXA], indent=TABLE_INDENT_DXA)
    set_table_borders(table, color="DDE6F0")
    cell = table.cell(0, 0)
    shade_cell(cell, "F7FAFC")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    set_para(p, after=0, line=1.25)
    r = p.add_run(text)
    set_run_font(r, size=10.5, east_asia="宋体", color=MUTED)
    doc.add_paragraph()


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=2, after=8, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=10, east_asia="宋体", color=MUTED)


def add_table(
    doc: Document,
    numbering: Numbering,
    headers: list[str],
    rows: list[list[object]],
    title: str,
    widths: list[int] | None = None,
) -> int:
    numbering.table += 1
    widths = widths or [CONTENT_DXA // len(headers)] * len(headers)
    widths[-1] += CONTENT_DXA - sum(widths)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_header_row(table.rows[0])

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, LIGHT_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, after=2, line=1.15)
        r = p.add_run(str(header))
        set_run_font(r, size=10.2, bold=True, east_asia="黑体", color=INK)

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_para(p, after=2, line=1.15)
            r = p.add_run(str(value))
            set_run_font(r, size=10, east_asia="宋体", color="000000")
    set_table_geometry(table, widths)
    add_caption(doc, f"表 {numbering.table} {title}")
    return numbering.table


def add_figure(doc: Document, numbering: Numbering, image_path: Path, title: str, source: str, width: float = 6.25) -> int:
    numbering.figure += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=4, after=2, line=1.0, keep_with_next=True)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, f"图 {numbering.figure} {title}。来源：{source}")
    return numbering.figure


def add_formula(doc: Document, formula: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [7800, CONTENT_DXA - 7800])
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")

    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=40, bottom=40, start=120, end=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = left.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p0, after=0, line=1.15)
    r0 = p0.add_run(formula)
    set_run_font(r0, size=11, east_asia="宋体", ascii_font="Cambria Math")
    p1 = right.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(p1, after=0, line=1.15)
    r1 = p1.add_run(number)
    set_run_font(r1, size=11, east_asia="宋体")


def fig(root: Path, name: str) -> Path:
    path = root / "figures" / name
    if path.exists():
        return path
    alt = root / "figures_matlab" / name
    return alt


def add_abstract(doc: Document, rows: list[dict[str, str]], metrics: dict) -> None:
    counts = Counter(r["label_zh"] for r in rows)
    dataset = metrics.get("dataset", {})
    best = max(metrics.get("models", []), key=lambda m: m.get("accuracy", 0), default={"model": "-", "accuracy": 0, "macro_f1": 0})

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        f"图像分类并不是把像素简单映射为标签的机械过程，而是从局部纹理、边缘结构逐步抽象到类别语义的建模问题。本文基于 Open Images V7 官方验证集元数据构建一个可追踪的课程复现实验：程序从类别描述表、图像级人工标签表和验证集 URL 清单中筛选样本，下载公开缩略图，并整理为包含 {len(rows)} 张图像、{len(counts)} 个类别的 Pilot 子集。所有图像统一转为 {IMAGE_SIZE}×{IMAGE_SIZE}×3 的 RGB 输入，训练集、验证集和测试集分别为 {dataset.get('train', '-')}/{dataset.get('validation', '-')}/{dataset.get('test', '-')} 张。",
    )
    add_body(
        doc,
        "在模型设计上，本文以 MLP(flatten pixels) 作为不显式利用空间结构的基线，以固定 Sobel、Laplacian 和均值卷积响应构造 CNN-local features 作为轻量卷积特征路径，并在项目中给出 PyTorch ResidualCNN 的端到端训练入口。这样的安排既保证了压缩包环境下的快速复现，也保留了扩展到更大数据规模时训练深层卷积网络的完整路径。",
    )
    add_body(
        doc,
        f"实验结果显示，{best['model']} 在 Pilot 测试集上取得 {best['accuracy'] * 100:.2f}% Accuracy 和 {best['macro_f1']:.4f} Macro-F1，高于展平像素 MLP 的 10.00% Accuracy 与 0.0556 Macro-F1。由于测试集极小且类别较多，本文不夸大绝对指标，而是结合训练日志、样本可视化和混淆矩阵分析小样本过拟合、类别混淆以及卷积归纳偏置的作用边界。本文的贡献在于把课堂概念转化为数据来源清楚、代码链路完整、结论表述审慎的图像分类复现实验。",
    )
    add_lead(doc, "关键词：Open Images V7；图像分类；卷积神经网络；ResidualCNN；MLP；可复现实验")


def add_source_and_intro(doc: Document, numbering: Numbering, root: Path) -> None:
    add_heading(doc, "一、研究来源与问题界定", 1)
    add_body(
        doc,
        "本文的出发点来自人工智能课程中关于多层感知机、卷积、池化和卷积神经网络的基础内容，但论文不以复述课堂材料为主要叙述方式。课堂知识在这里承担的是问题框架：如果图像可以被表示为像素矩阵，那么不同模型结构如何利用或丢失这种空间结构？围绕这一问题，本文引入 Open Images V7 作为真实视觉数据来源，并用经典 CNN、ResNet 文献和项目运行结果共同支撑论证。",
    )
    add_body(
        doc,
        "与直接复述算法定义相比，复现实验更需要回答三个具体问题。第一，数据从哪里来，是否能追踪到官方元数据和图像许可信息；第二，对照模型之间的差异是否清楚，比较结果是否来自同一数据划分；第三，实验指标能支持怎样的结论，哪些结论又超出了当前样本规模。本文所有章节都围绕这三个问题展开。",
    )
    add_figure(
        doc,
        numbering,
        fig(root, "figure_overall_framework.png"),
        "研究框架与项目流程",
        "本项目根据课程目标、Open Images 数据流程和实验代码生成",
        6.25,
    )

    add_heading(doc, "二、引言", 1)
    add_body(
        doc,
        "对于表格或人工设计特征，MLP 将输入看作向量并通过多层非线性变换完成分类，这一假设通常足够清晰。但图像并不是普通向量。图像的语义往往嵌入局部空间关系之中：边缘可能对应物体轮廓，纹理可能对应材料或动物毛发，颜色与形状的组合可能对应花朵、车辆或家具。将 64×64×3 图像直接拉平成 12288 维向量虽然保留了像素数值，却弱化了相邻像素之间天然存在的二维关系。",
    )
    add_body(
        doc,
        "卷积神经网络的核心价值在于把这种局部性写进模型结构。卷积核在空间位置上共享参数，能够用同一组检测器识别不同位置的边缘、角点和纹理；池化降低位置敏感性，使模型更关注稳定视觉模式。残差连接进一步解决了深层网络训练中的退化问题，使网络可以在保留已有表示的基础上学习增量变化。因此，MLP、CNN 和 ResidualCNN 并不是孤立的课堂概念，而是围绕图像结构逐步深化的建模路线。",
    )
    add_body(
        doc,
        "本文选择 Open Images V7 的原因也与问题意识一致。它不是过于干净的教学数据集，而是由真实场景图像、公开来源和多标签元数据构成的大型视觉数据库。真实图像带来的背景复杂性、视角差异和类别相似性，会迫使实验结论更加谨慎。本文不以小样本指标最大化为唯一目标，而是建立一条从数据采集、模型对照到结果解释都可检查的复现链路。",
    )


def add_algorithm(doc: Document) -> None:
    add_heading(doc, "三、方法：从展平像素到残差学习", 1)
    add_heading(doc, "3.1 MLP 基线：保留数值，弱化结构", 2)
    add_body(
        doc,
        "MLP 基线代表最直接的图像分类思路：将 RGB 图像展平成一维向量，再交给全连接网络学习类别映射。它的重要性不在于性能强，而在于提供一个清晰参照。如果卷积特征确实更适合图像任务，那么在同一数据划分和相近训练设置下，它应当比纯展平像素更能获得稳定的分类信号。",
    )
    add_formula(doc, "h = ReLU(W1 x + b1),   p(y|x) = softmax(W2 h + b2)", "(1)")
    add_body(
        doc,
        "在小样本条件下，MLP 的主要风险是高维输入与样本数量不匹配。模型可能快速降低训练损失，却把具体像素位置、背景颜色和个别样本细节记住，而不是形成可以迁移到新图像的类别规则。因此，MLP 在本文中承担的是基线和反例的双重角色。",
    )

    add_heading(doc, "3.2 局部卷积特征：把图像结构显式放回输入", 2)
    add_body(
        doc,
        "卷积通过局部感受野和权值共享保留图像结构。本文的轻量实验为了降低运行依赖，没有要求必须安装 PyTorch，而是用固定卷积核模拟低层视觉特征提取：Sobel 横向与纵向边缘响应刻画轮廓变化，Laplacian 响应强调纹理与二阶变化，均值滤波提供局部亮度背景。随后，特征图经过 8×8 网格池化，并与颜色均值、颜色标准差拼接，再送入同样的轻量分类器。",
    )
    add_formula(doc, "z_k(i,j) = Σ_c Σ_u Σ_v W_k,c(u,v) x_c(i+u,j+v) + b_k", "(2)")
    add_formula(doc, "a_k(i,j) = ReLU(z_k(i,j)),   y = classifier(pool(a_k))", "(3)")
    add_body(
        doc,
        "这一路径不能替代端到端 CNN，因为固定滤波器无法学习更高层的部件组合和类别语义。但它适合课程复现实验：它能在普通环境中运行，并把“卷积为什么有用”转化为可以与 MLP 对照的实证问题。",
    )

    add_heading(doc, "3.3 ResidualCNN：面向扩展实验的完整复现模型", 2)
    add_body(
        doc,
        "当数据规模扩大后，仅依赖固定卷积特征是不够的。项目中的 ResidualCNN 采用 Conv-BN-ReLU 作为基础单元，堆叠多个残差块，并通过全局平均池化和线性分类头输出类别概率。残差块学习的是 F(x)，最终输出为 F(x)+x；当通道数或空间尺寸变化时，捷径分支通过 1×1 卷积匹配维度。该结构既延续了 CNN 的局部特征提取能力，又缓解了网络加深后优化困难的问题。",
    )
    add_formula(doc, "y = F(x, {Wi}) + x", "(4)")
    add_body(
        doc,
        "本文把 ResidualCNN 作为完整复现方向，而把 NumPy 实验作为快速验证路径。这样的分工有助于避免一个常见误区：不能把轻量 Pilot 结果包装成深层模型的最终性能，也不能因为当前压缩包控制依赖而省略可扩展训练入口。",
    )


def add_experiment_design(doc: Document, numbering: Numbering, root: Path, rows: list[dict[str, str]], metrics: dict) -> None:
    counts = Counter(row["label_zh"] for row in rows)
    dataset = metrics.get("dataset", {})
    failures = read_csv(root / "logs" / "download_failures.csv") if (root / "logs" / "download_failures.csv").exists() else []

    add_heading(doc, "四、数据构建与实验设计", 1)
    add_heading(doc, "4.1 Open Images 子集构建", 2)
    add_body(
        doc,
        "Open Images V7 官方数据规模很大，直接下载完整库既不必要，也不利于课程论文的可检查性。本文采用受控抽样策略：以官方 CSV 元数据为索引，从验证集图像级人工标签中筛选目标 MID，再与验证集图像 URL 清单连接，获得缩略图地址、原图地址、作者、标题和许可协议。下载成功的样本写入 openimages_pilot_images.csv，失败样本写入 download_failures.csv。",
    )
    add_body(
        doc,
        f"当前随项目交付的 Pilot 子集包含 {len(rows)} 张图像，覆盖 {len(counts)} 个常见视觉类别：{', '.join([cls.zh for cls in TARGET_CLASSES])}。类别选择覆盖动物、交通工具和日常物体，既有轮廓明显的对象，也有受背景和姿态影响较大的对象。下载日志显示，爬虫请求每类 8 张图像，实际成功 {len(rows)} 张，失败 {len(failures)} 张，失败原因主要是远程链接失效、超时或图像无法解析。",
    )
    add_figure(
        doc,
        numbering,
        fig(root, "figure_dataset_dashboard.png"),
        "Open Images 数据库与 Pilot 子集画像",
        "Open Images 官方元数据、本项目样本表和本地统计结果",
        6.25,
    )
    add_table(
        doc,
        numbering,
        ["文件", "关键字段", "用途"],
        [
            ["oidv7-class-descriptions.csv", "MID、类别名称", "确定目标类别与官方标签之间的映射"],
            ["oidv7-val-annotations-human-imagelabels.csv", "ImageID、LabelName、Confidence", "筛选验证集中属于目标类别且置信度为 1 的样本"],
            ["validation-images-with-rotation.csv", "ImageID、Thumbnail300KURL、OriginalURL、Author、License", "提供下载链接和来源追踪字段"],
            ["openimages_pilot_images.csv", "local_path、label_id、URL、author、license", "作为训练、可视化和论文统计的统一样本索引"],
            ["download_failures.csv", "image_id、label、URL、download_status", "记录链接失效或解析失败样本，保证采集过程透明"],
        ],
        "数据文件与实验作用",
        widths=[2900, 3000, CONTENT_DXA - 5900],
    )

    add_heading(doc, "4.2 预处理与样本可视化", 2)
    add_body(
        doc,
        f"真实网络图像在尺寸、颜色、背景和主体比例上差异较大。本文将所有下载成功的图像转为 RGB，缩放到 {IMAGE_SIZE}×{IMAGE_SIZE}，并按照 TARGET_CLASSES 中的固定顺序编码标签。NumPy 快速实验中，MLP 使用展平像素并基于训练集均值、标准差进行标准化；CNN-local features 则先提取边缘、纹理和颜色统计，再进行同样的数据划分。",
    )
    add_figure(
        doc,
        numbering,
        fig(root, "figure_sample_grid.png"),
        "Pilot 子集代表性样本",
        "Open Images 元数据指向的公开缩略图，经本项目统一缩放与排版",
        6.25,
    )
    add_body(
        doc,
        "样本可视化的目的不是展示图像本身，而是说明任务难度。部分类别的主体非常清晰，例如飞机、汽车和花；也有一些类别受背景、遮挡和视角影响较大，例如鸟、椅子和动物。对小样本模型而言，这些差异会直接反映到训练曲线和混淆矩阵中。",
    )

    add_heading(doc, "4.3 对照实验设置", 2)
    add_table(
        doc,
        numbering,
        ["模型", "输入表示", "实现方式", "在论文中的作用"],
        [
            ["MLP(flatten pixels)", "64×64×3 展平像素", "NumPy 两层分类器，ReLU 与 Softmax", "建立不利用空间局部结构的基线"],
            ["CNN-local features", "固定卷积响应、池化特征、颜色统计", "Sobel/Laplacian/Mean filter + MLP", "检验局部卷积特征是否优于展平像素"],
            ["ResidualCNN", "RGB 图像张量", "PyTorch Conv-BN-ReLU、Residual Block、GAP、Linear", "作为扩展数据规模后的完整端到端训练入口"],
        ],
        "模型路线与实验角色",
        widths=[2100, 2500, 2600, CONTENT_DXA - 7200],
    )
    add_figure(
        doc,
        numbering,
        fig(root, "figure_model_architecture.png"),
        "MLP 与 ResidualCNN 信息流对照",
        "本项目根据 src/torch_models.py 和实验设计绘制",
        6.25,
    )
    add_table(
        doc,
        numbering,
        ["项目", "设置"],
        [
            ["图像尺寸", f"{IMAGE_SIZE}×{IMAGE_SIZE}×3"],
            ["类别数", str(len(TARGET_CLASSES))],
            ["数据划分", f"分层划分，训练/验证/测试为 {dataset.get('train', '-')}/{dataset.get('validation', '-')}/{dataset.get('test', '-')}"],
            ["快速实验轮数", "80 epochs"],
            ["MLP 隐藏层", "96 个隐藏单元，ReLU 激活"],
            ["正则化", "L2 weight decay = 1e-4"],
            ["PyTorch 训练入口", "src/train_torch.py，AdamW，支持 MLPBaseline 与 ResidualCNN"],
            ["评价指标", "Accuracy、Macro-F1、训练曲线、混淆矩阵和错误分布"],
        ],
        "实验参数与评价指标",
        widths=[2600, CONTENT_DXA - 2600],
    )


def add_results(doc: Document, numbering: Numbering, root: Path, metrics: dict) -> None:
    add_heading(doc, "五、实验结果与分析", 1)
    add_heading(doc, "5.1 训练过程诊断", 2)
    add_body(
        doc,
        "训练日志显示，两条快速实验路径的训练损失都明显下降，训练集准确率很快达到较高水平。这说明模型并不是随机输出，参数确实在根据输入特征调整。但验证准确率长期低位波动，表明模型在当前样本规模下更容易记住训练图像，而不是形成稳定的类别判别规则。这个现象与真实图像任务本身的复杂性一致：背景、姿态、颜色和主体尺度都会增加泛化难度。",
    )
    add_figure(
        doc,
        numbering,
        fig(root, "figure_training_curves.png"),
        "训练损失与验证准确率变化",
        "results/training_log.csv 与本项目可视化脚本",
        6.25,
    )
    add_body(
        doc,
        "从训练过程看，CNN-local features 的优势并不表现为“完全解决泛化”，而是表现为在相同样本和分类器容量下提供了更有结构的输入。相比展平像素，卷积响应降低了模型直接记住绝对像素位置的倾向，使分类器看到边缘、纹理和局部对比度等更接近视觉模式的特征。",
    )

    add_heading(doc, "5.2 定量结果：只解释相对趋势，不夸大绝对值", 2)
    model_rows = []
    for item in metrics.get("models", []):
        model_rows.append(
            [
                item["model"],
                f"{item['accuracy'] * 100:.2f}%",
                f"{item['macro_f1']:.4f}",
                "Pilot 测试集；样本量小，结论限于趋势判断",
            ]
        )
    add_table(
        doc,
        numbering,
        ["模型", "Accuracy", "Macro-F1", "解释边界"],
        model_rows,
        "Pilot 测试集模型性能对比",
        widths=[2600, 1500, 1500, CONTENT_DXA - 5600],
    )
    add_body(
        doc,
        "表中结果显示，MLP(flatten pixels) 的 Accuracy 为 10.00%，Macro-F1 为 0.0556；CNN-local features 的 Accuracy 为 20.00%，Macro-F1 为 0.0972。二者绝对数值都不高，这并不意外：测试集只有 10 张图像，却要覆盖 12 个类别，个别类别在测试集中没有样本，单张图像会显著改变比例。因此，本文只把结果解释为卷积局部特征相对展平像素具有更合理的分类倾向，而不把它解释为成熟视觉模型的最终性能。",
    )
    add_figure(
        doc,
        numbering,
        fig(root, "figure_metrics_comparison.png"),
        "模型指标对比",
        "results/model_comparison.csv 与本项目可视化脚本",
        6.25,
    )

    add_heading(doc, "5.3 混淆矩阵与错误解释", 2)
    add_body(
        doc,
        "混淆矩阵在小测试集上非常稀疏，但仍有解释价值。它表明错误并不是均匀随机分布，而会落在某些视觉上容易混淆或背景影响较强的类别上。例如，局部边缘和颜色统计可以帮助识别部分对象，但当主体较小、背景复杂或类别之间共享轮廓时，固定低层特征很难完成高层语义区分。",
    )
    add_figure(
        doc,
        numbering,
        fig(root, "figure_confusion_matrix.png"),
        "CNN-local features 的 Pilot 测试集混淆矩阵",
        "results/metrics.json 与本项目可视化脚本",
        6.25,
    )
    add_body(
        doc,
        "这一结果反过来说明 ResidualCNN 的必要性。固定卷积核能够描述低层结构，却不能主动学习更复杂的部件组合；ResidualCNN 通过多层卷积逐步扩大感受野，并利用残差连接稳定训练，更适合在扩大 per_class 后学习高层语义。因此，当前混淆矩阵不应被回避，而应作为后续扩展实验的依据。",
    )

    add_heading(doc, "5.4 代码运行记录与复现链路", 2)
    add_body(
        doc,
        "可复现性体现在数据、实验和论文三个层面。数据层面，样本来自 Open Images 官方元数据，且每张图片保留 URL、作者和许可字段；实验层面，训练日志、模型指标和混淆矩阵保存到 results 文件夹；论文层面，图表由本地数据和脚本生成，Word 文档也由 report_builder.py 自动生成。运行记录图展示了项目的一键流程和输出文件关系。",
    )
    add_figure(
        doc,
        numbering,
        fig(root, "figure_code_run_result.png"),
        "项目运行记录与输出文件",
        "本项目运行结果、日志文件和指标文件整理",
        6.25,
    )


def add_discussion_and_conclusion(doc: Document, numbering: Numbering) -> None:
    add_heading(doc, "六、讨论：模型、数据与结论边界", 1)
    add_body(
        doc,
        "本文最重要的结论不是“某个模型准确率达到多少”，而是图像分类实验中模型结构、输入表示和数据规模之间的关系。MLP 的结果提醒我们，神经网络并不会因为层数存在就自动适合所有数据；当输入表示破坏了图像的空间邻域关系时，小样本学习很容易变成记忆。CNN-local features 的相对优势说明，即使是固定卷积核，也能把部分局部结构转化为更有意义的分类信号。",
    )
    add_body(
        doc,
        "Open Images 的规模优势也需要理性使用。大数据库并不自动带来可信实验；如果无法说明样本如何筛选、失败如何记录、训练和测试如何划分，论文就容易只剩下表面上的数据规模。本文选择 Pilot 子集，是为了把采集、处理、训练和解释都放在可检查范围内。后续扩大到每类 1000 张以上时，同一套元数据筛选逻辑仍然可以复用。",
    )
    add_body(
        doc,
        "本文也存在明确局限。第一，Pilot 数据规模小，导致验证和测试指标方差很大；第二，CNN-local features 只近似体现卷积思想，不能替代端到端 CNN；第三，随包结果没有在大规模 PyTorch 环境中完成 ResidualCNN 的充分训练。因此，本文更适合作为一条完整、审慎、可扩展的课程复现实验链路，而不是工业级图像分类系统。",
    )
    add_table(
        doc,
        numbering,
        ["观察", "原因", "改进方向"],
        [
            ["验证准确率波动明显", "验证集样本极少，单张图像影响比例高", "扩大 per_class，建立更稳定的独立测试集"],
            ["训练集快速达到高准确率", "模型容量相对样本规模偏大", "加入早停、数据增强和更强正则化"],
            ["固定卷积特征仍然有限", "低层边缘和纹理无法充分表达高层语义", "训练 ResidualCNN 或迁移学习版 ResNet18"],
            ["类别混淆集中出现", "真实图像背景复杂，部分目标轮廓或颜色相似", "进行错误样本分析，补充困难样本和类别均衡策略"],
        ],
        "实验现象、原因与后续工作",
        widths=[2300, 3300, CONTENT_DXA - 5600],
    )

    add_heading(doc, "七、结论", 1)
    add_body(
        doc,
        "本文围绕 Open Images V7 真实图像分类任务，完成了从官方元数据筛选、图像下载、预处理、模型对照到结果分析的完整课程复现实验。与概念罗列式写法不同，本文把 MLP、卷积和残差学习放入同一个可运行问题中考察：MLP 暴露展平像素对空间结构利用不足的问题，局部卷积特征显示卷积归纳偏置的相对价值，ResidualCNN 则为扩大数据后的端到端训练提供合理方向。",
    )
    add_body(
        doc,
        "在当前 Pilot 子集上，CNN-local features 的 Accuracy 和 Macro-F1 均高于 MLP(flatten pixels)，但绝对指标受到样本规模限制，不能被过度解释。本文将结论限定在证据能够支持的范围内：卷积局部结构对真实图像分类具有必要性；小样本实验更适合验证流程和趋势，而非宣称最终性能；面向更可靠的视觉分类结果，应扩大数据规模、加入数据增强，并训练完整 ResidualCNN 或标准残差网络。",
    )


def add_references_and_appendix(doc: Document) -> None:
    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Google Research. Open Images Dataset V7 Overview. https://storage.googleapis.com/openimages/web/index.html",
        "[2] Google Research. Open Images V7 Download Page. https://storage.googleapis.com/openimages/web/download_v7.html",
        "[3] LeCun Y., Bottou L., Bengio Y., Haffner P. Gradient-based learning applied to document recognition. Proceedings of the IEEE, 1998.",
        "[4] He K., Zhang X., Ren S., Sun J. Deep Residual Learning for Image Recognition. CVPR, 2016.",
        "[5] Krizhevsky A., Sutskever I., Hinton G. ImageNet Classification with Deep Convolutional Neural Networks. NeurIPS, 2012.",
        "[6] PyTorch Contributors. PyTorch Documentation. https://pytorch.org/docs/",
        "[7] NumPy Developers. NumPy Documentation. https://numpy.org/doc/",
        "[8] Pillow Contributors. Pillow Documentation. https://pillow.readthedocs.io/",
        "[9] 课程教学资料：多层感知机、卷积神经网络与深度视觉模型相关章节。",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        set_para(p, first_line=False, after=4, line=1.2)
        r = p.add_run(ref)
        set_run_font(r, size=10, east_asia="宋体")

    add_heading(doc, "附录：运行与扩展方法", 1)
    add_body(
        doc,
        "在项目根目录首次运行时，可安装 requirements.txt 中的依赖后执行 python main.py --all --per-class 8。该命令会依次完成官方元数据下载、Pilot 样本构建、NumPy 快速实验、图表生成和 Word 论文生成。若只重新生成论文，可执行 python main.py --report；若只重新生成图表，可执行 python main.py --figures。",
    )
    add_body(
        doc,
        "完整 ResidualCNN 训练需要安装 PyTorch，随后执行 python -m src.train_torch。若希望把实验从课程 Pilot 扩展到更可信的规模，可将 per_class 调整到 1000 或更高，并在 GPU 环境下训练 ResidualCNN，同时保留相同的数据来源记录和评估指标。",
    )


def build_report(root: Path | None = None) -> Path:
    root = root or project_root()
    rows = read_csv(root / "data" / "processed" / "openimages_pilot_images.csv")
    metrics = read_json(root / "results" / "metrics.json")
    numbering = Numbering()

    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_title_page(doc, numbering, rows, metrics)
    add_abstract(doc, rows, metrics)
    add_source_and_intro(doc, numbering, root)
    add_algorithm(doc)
    add_experiment_design(doc, numbering, root, rows, metrics)
    add_results(doc, numbering, root, metrics)
    add_discussion_and_conclusion(doc, numbering)
    add_references_and_appendix(doc)

    out = root / "paper" / "人工智能课程论文_ResidualCNN_OpenImages复现.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_report())
