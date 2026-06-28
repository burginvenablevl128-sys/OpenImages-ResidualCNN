from __future__ import annotations

import csv
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "paper" / "人工智能课程论文复现_完整结构版.docx"

INK = "111827"
BLUE = "1F4D78"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F4F6"
MUTED = "4B5563"
BORDER = "CBD5E1"
CONTENT_DXA = 9072
TABLE_INDENT_DXA = 0


class Numbering:
    def __init__(self) -> None:
        self.fig = 0
        self.tbl = 0


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_run_font(
    run,
    size: float = 12,
    bold: bool = False,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
    color: str = "000000",
    italic: bool = False,
) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def set_para(p, before=0, after=6, line=1.5, first_line=False, keep=False) -> None:
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.3)
    section.footer_distance = Cm(1.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    for style_name, size, before, after in [
        ("Heading 1", 16, 14, 8),
        ("Heading 2", 14, 10, 6),
        ("Heading 3", 12.5, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    run._r.append(text)
    run._r.append(end)


def header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(header, after=0, line=1)
    set_run_font(header.add_run("Open Images 图像分类复现实验"), 9, east_asia="宋体", color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(footer, after=0, line=1)
    set_run_font(footer.add_run("第 "), 9, east_asia="宋体", color=MUTED)
    page_number(footer)
    set_run_font(footer.add_run(" 页"), 9, east_asia="宋体", color=MUTED)
    for run in footer.runs:
        set_run_font(run, 9, east_asia="宋体", color=MUTED)


def h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(p, before=14 if level == 1 else 10, after=8 if level == 1 else 5, line=1.25, keep=True)
    for run in p.runs:
        set_run_font(run, 16 if level == 1 else 14 if level == 2 else 12.5, True, "黑体", color=BLUE)


def body(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    set_para(p, first_line=first_line)
    set_run_font(p.add_run(text), 12, east_asia="宋体", color="000000")


def compact(doc: Document, text: str, color: str = MUTED) -> None:
    p = doc.add_paragraph()
    set_para(p, before=0, after=4, line=1.25, first_line=False)
    set_run_font(p.add_run(text), 10.5, east_asia="宋体", color=color)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=2, after=8, line=1.2)
    set_run_font(p.add_run(text), 10.5, east_asia="宋体", color=MUTED)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(table, color=BORDER, val="single") -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tbl_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tbl_borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, width in enumerate(widths):
            cell = row.cells[i]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def table_doc(doc: Document, num: Numbering, headers: list[str], rows: list[list[object]], title: str, widths: list[int]) -> None:
    num.tbl += 1
    widths[-1] += CONTENT_DXA - sum(widths)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    geometry(table, widths)
    borders(table)
    repeat_header(table.rows[0])

    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=100, bottom=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, after=0, line=1.15)
        set_run_font(p.add_run(value), 10.5, True, "黑体", color=INK)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=80, bottom=80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(str(value)) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            set_para(p, after=0, line=1.15)
            set_run_font(p.add_run(str(value)), 10.5, east_asia="宋体", color="000000")
    geometry(table, widths)
    caption(doc, f"表{num.tbl} {title}")


def formula(doc: Document, expr: str, no: str, explain: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    geometry(table, [7400, CONTENT_DXA - 7400])
    borders(table, val="nil")
    left, right = table.rows[0].cells
    for cell in (left, right):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=40, bottom=40)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=0, line=1.2)
    set_run_font(p.add_run(expr), 11.5, east_asia="宋体", ascii_font="Cambria Math")
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(p, after=0, line=1.2)
    set_run_font(p.add_run(no), 11.5)
    if explain:
        body(doc, explain)


def fig_path(name: str) -> Path:
    return ROOT / "figures" / name


def figure(doc: Document, num: Numbering, path: Path, title: str, source: str, width=6.05) -> None:
    if not path.exists():
        return
    num.fig += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=4, after=2, line=1, keep=True)
    p.add_run().add_picture(str(path), width=Inches(width))
    caption(doc, f"图{num.fig} {title}。来源：{source}")


def title_page(doc: Document, num: Numbering, rows: list[dict[str, str]], metrics: dict) -> None:
    dataset = metrics.get("dataset", {})
    counts = Counter(r["label_zh"] for r in rows)
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=18, line=1.2)
    set_run_font(p.add_run("《人工智能》课程论文"), 16, True, "黑体", color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=10, line=1.2)
    set_run_font(p.add_run("从像素到语义：基于 Open Images 子集的残差卷积神经网络图像分类复现实验"), 22, True, "黑体", color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=24, line=1.2)
    set_run_font(p.add_run("以 MLP 基线为参照的 CNN 结构价值分析"), 14, False, "宋体", color=MUTED)

    table_doc(
        doc,
        num,
        ["项目", "内容"],
        [
            ["课程论文方向", "人工智能算法复现与应用分析"],
            ["数据来源", "Open Images V7 官方公开视觉数据库"],
            ["复现算法", "卷积神经网络思想与 ResidualCNN 训练框架"],
            ["对照模型", "MLP(flatten pixels)、CNN-local features"],
            ["实验规模", f"{len(rows)} 张图像，{len(counts)} 个类别，训练/验证/测试为 {dataset.get('train', '-')}/{dataset.get('validation', '-')}/{dataset.get('test', '-')}"],
        ],
        "论文基本信息",
        [2100, CONTENT_DXA - 2100],
    )

    compact(doc, "正文采用小四号中文字体、1.5 倍行距；表格文字比正文略小，英文和数字采用 Times New Roman；公式置于无边框表格中并右侧编号。")
    doc.add_page_break()


def source_statement(doc: Document, num: Numbering) -> None:
    h(doc, "一、课程论文来源说明", 1)
    body(doc, "本文选题为“基于 Open Images 子集的残差卷积神经网络图像分类复现实验”。论文以人工智能课程中多层感知机、卷积、池化、损失函数和模型评估等内容为基础，进一步围绕真实图像分类任务完成数据构建、算法复现、对照实验和结果分析。")
    body(doc, "论文的数据来源为 Google Research 发布的 Open Images V7。本文没有直接下载完整数据库，而是基于官方类别描述表、图像级人工标签表和验证集图像 URL 表筛选 12 个常见视觉类别，下载公开缩略图并整理为 Pilot 子集。每张图片均在 openimages_pilot_images.csv 中保留 image_id、类别、URL、作者、许可协议和本地路径，便于追溯。")
    body(doc, "实验代码由本人围绕课程论文任务完成和整理，包括 Open Images 元数据下载、样本筛选、图像预处理、NumPy 快速实验、PyTorch ResidualCNN 入口、训练日志保存、指标计算和论文图表生成。论文参考了 Open Images 官方资料、LeNet、AlexNet、ResNet 相关论文，以及 PyTorch、NumPy、Pillow 等官方文档；未直接复制网络博客或开源项目代码。")
    body(doc, "论文写作和排版过程中使用了大语言模型辅助，主要用于课程论文结构建议、概念解释、文字润色、段落衔接、Word 版式整理和缺失章节补写。实验数据、运行结果、图表文件和样本来源均来自本地项目文件；模型结果没有由大语言模型虚构。本人独立完成实验主题选择、代码运行材料整理、结果核对与最终提交判断。")
    table_doc(
        doc,
        num,
        ["资料类别", "具体来源", "在论文中的作用"],
        [
            ["课程材料", "人工智能课程 PPT 与课堂讲授内容", "提供 MLP、CNN、池化、损失函数和评价指标的理论基础"],
            ["公开数据", "Open Images V7 官方网页与下载说明", "提供图像、标签、URL、作者和许可信息"],
            ["经典论文", "LeNet、AlexNet、ResNet 相关论文", "支撑卷积结构与残差连接的理论说明"],
            ["软件文档", "PyTorch、NumPy、Pillow 官方文档", "支撑模型实现、数值计算和图像处理"],
            ["本地项目", "data、results、figures、logs、paper 文件夹", "提供样本表、训练日志、评估指标、图表和运行记录"],
            ["大语言模型", "Codex/ChatGPT 辅助", "用于结构建议、文字润色、概念解释和排版生成，不替代实验结果"],
        ],
        "课程论文资料来源与使用方式",
        [1700, 3300, CONTENT_DXA - 5000],
    )


def abstract_intro(doc: Document, num: Numbering, rows: list[dict[str, str]], metrics: dict) -> None:
    dataset = metrics.get("dataset", {})
    best = max(metrics.get("models", []), key=lambda m: m.get("accuracy", 0))
    h(doc, "二、摘要", 1)
    body(doc, f"图像分类看似只是给图片贴标签，背后却包含从像素到语义的建模问题。若把图像直接展平成一维向量，模型虽然可以接收输入，却很难自然利用相邻像素之间的边缘、纹理和局部形状关系；卷积神经网络的价值正是在于把这种空间局部性写入模型结构。本文基于 Open Images V7 公开视觉数据库，构建包含 {len(rows)} 张真实图像、12 个类别的 Pilot 图像分类子集，围绕 MLP 基线、局部卷积特征模型和 ResidualCNN 训练框架完成一次可复现的课程实验。")
    body(doc, f"实验从数据采集开始：程序下载 Open Images 官方类别表、图像级人工标签表和验证集 URL 表，根据目标类别 MID 筛选样本，下载缩略图并统一转为 64×64×3 的 RGB 图像。当前数据划分为训练集 {dataset.get('train', '-')} 张、验证集 {dataset.get('validation', '-')} 张、测试集 {dataset.get('test', '-')} 张。模型设计上，本文先用 MLP(flatten pixels) 作为不利用空间结构的基线，再使用 Sobel、Laplacian 和均值滤波构造 CNN-local features，最后给出可扩展的 PyTorch ResidualCNN 端到端训练代码。")
    body(doc, f"当前可运行结果显示，{best['model']} 在 Pilot 测试集上的准确率为 {best['accuracy'] * 100:.2f}%，Macro-F1 为 {best['macro_f1']:.4f}，高于展平像素 MLP。由于 Pilot 数据规模有意控制在适合课程检查的范围内，本文没有把这一数值解释为模型最终性能，而是结合训练曲线、混淆矩阵和错误案例讨论小样本过拟合、类别混淆和数据扩展路径。论文最终形成了从公开数据库采集、数据预处理、模型搭建、训练记录、实验结果到学术化讨论的完整闭环。")
    compact(doc, "关键词：Open Images；图像分类；卷积神经网络；ResidualCNN；MLP；算法复现", color=INK)

    h(doc, "三、引言", 1)
    body(doc, "人工智能在图像识别、自动驾驶、医疗影像、工业检测和内容检索等场景中已经成为重要技术基础。课程中学习的神经网络并不是抽象公式，而是能够被用于真实数据任务的建模工具。图像分类任务之所以适合作为课程论文复现对象，是因为它能够把输入表示、网络结构、损失函数、训练过程和评价指标连接在同一个实验链路中。")
    body(doc, "在多层感知机中，图像常被展平成向量后输入全连接层。这种方式实现简单，却忽略了图像本来的二维空间关系。相邻像素往往共同形成边缘、纹理、轮廓和局部形状，而这些局部结构正是视觉语义的重要来源。卷积神经网络通过局部感受野、权值共享和池化操作，使模型能够更自然地提取图像局部模式。")
    body(doc, "本文选择 ResidualCNN 作为完整复现方向，是因为残差连接代表现代深层视觉网络中的关键思想。普通 CNN 能够提取层级特征，但网络加深后可能出现训练退化；残差结构通过捷径连接缓解深层优化困难，使模型在扩大数据集后具有更好的可训练性。本文同时保留 MLP 和 CNN-local features 对照实验，是为了让“卷积结构为什么有价值”这个问题在同一数据集上得到直观回答。")
    figure(doc, num, fig_path("figure_overall_framework.png"), "研究框架与项目流程", "本项目根据课程要求、数据流程和模型复现实验用 Python/PIL 自行绘制")


def algorithm(doc: Document) -> None:
    h(doc, "四、算法简介", 1)
    h(doc, "4.1 MLP 基线模型", 2)
    body(doc, "MLP 基线代表最直接的图像分类思路：将 64×64×3 的 RGB 图像展平为 12288 维向量，然后输入全连接网络完成类别预测。该模型的优点是结构简单、计算路径清楚，适合作为不利用空间结构的参照模型；缺点是参数与像素位置强绑定，很难在小样本条件下稳定学习局部形状和纹理。")
    formula(doc, "h = ReLU(W_1x + b_1),  p = softmax(W_2h + b_2)", "(4-1)", "其中，x 为展平后的像素向量，h 为隐藏层表示，p 为模型输出的类别概率分布。")

    h(doc, "4.2 卷积、池化与局部特征", 2)
    body(doc, "卷积层的核心思想是局部连接和权值共享。一个卷积核只观察输入图像中的局部窗口，却会在整张图像上重复使用，因此同一个边缘或纹理检测器可以在不同位置生效。池化操作进一步压缩空间尺寸，降低模型对像素精确位置的敏感性。")
    formula(doc, "z_k(i,j) = Σ_c Σ_u Σ_v W_{k,c}(u,v)x_c(i+u,j+v) + b_k", "(4-2)", "其中，W 表示卷积核参数，x 表示输入图像，z_k(i,j) 表示第 k 个卷积核在位置 (i,j) 的响应。")
    formula(doc, "a_k(i,j) = ReLU(z_k(i,j)) = max(0, z_k(i,j))", "(4-3)", "ReLU 激活函数用于引入非线性，使模型能够表达更复杂的视觉模式。")

    h(doc, "4.3 残差学习", 2)
    body(doc, "当卷积网络变深时，理论表达能力增强，但训练过程不一定更容易。ResNet 的核心思想是让若干卷积层学习残差函数 F(x)，再与输入 x 相加得到输出。若某些层暂时无法学习有效变化，捷径连接至少能够保留输入信息，从而减轻深层网络退化问题。")
    formula(doc, "y = F(x, {W_i}) + x", "(4-4)", "其中，F(x,{W_i}) 表示残差分支学习到的变换，x 表示捷径连接传递的输入特征。")

    h(doc, "4.4 损失函数与评价指标", 2)
    body(doc, "本文采用交叉熵损失衡量真实标签与预测概率之间的差异，并在快速实验中加入 L2 正则项缓解小样本过拟合。分类结果使用 Accuracy、Precision、Recall、F1、Macro-F1 和混淆矩阵进行评价，其中 Macro-F1 能避免整体准确率掩盖少数类别表现。")
    formula(doc, "L = -Σ_c y_c log(p_c) + λ||θ||_2^2", "(4-5)", "其中，y_c 为真实标签的 one-hot 编码，p_c 为预测概率，λ||θ||_2^2 为正则项。")
    formula(doc, "Precision = TP/(TP+FP),  Recall = TP/(TP+FN),  F1 = 2PR/(P+R)", "(4-6)")


def experiment(doc: Document, num: Numbering, rows: list[dict[str, str]], metrics: dict) -> None:
    counts = Counter(r["label_zh"] for r in rows)
    dataset = metrics.get("dataset", {})
    failures = read_csv(ROOT / "logs" / "download_failures.csv") if (ROOT / "logs" / "download_failures.csv").exists() else []

    h(doc, "五、实验设计与应用场景", 1)
    body(doc, "本文实验面向真实网络图像的多类别分类场景，可类比应用于简单的图像检索、内容归档、教学演示和低资源视觉识别任务。实验设计强调“可复现”和“可解释”：数据要能追溯来源，预处理要能重复，模型对照要控制变量，结果分析要说明边界。")
    h(doc, "5.1 数据介绍", 2)
    h(doc, "（一）数据集介绍", 3)
    body(doc, "Open Images V7 是 Google Research 发布的大规模公开视觉数据库，包含大量真实场景图像和图像级标签。本文不追求完整下载，而是将其作为数据母体，筛选猫、狗、汽车、自行车、摩托车、鸟、马、船、飞机、椅子、卡车、花 12 个常见类别，构建课程可快速检查的 Pilot 子集。")
    table_doc(
        doc,
        num,
        ["项目", "数值或说明"],
        [
            ["数据母体", "Open Images V7 官方公开视觉数据库"],
            ["实验子集", f"{len(rows)} 张图像，12 个类别"],
            ["类别分布", "、".join(f"{k}{v}张" for k, v in counts.items())],
            ["图像尺寸", "统一为 64×64×3 RGB"],
            ["数据划分", f"训练集 {dataset.get('train', '-')} 张，验证集 {dataset.get('validation', '-')} 张，测试集 {dataset.get('test', '-')} 张"],
            ["失败记录", f"下载失败或解析失败样本 {len(failures)} 条，已记录到 logs/download_failures.csv"],
        ],
        "Open Images Pilot 子集统计",
        [2300, CONTENT_DXA - 2300],
    )
    figure(doc, num, fig_path("figure_dataset_dashboard.png"), "Open Images 数据库与 Pilot 子集概览", "Open Images 官方元数据与本项目样本表，Python/PIL 自行绘制")

    h(doc, "（二）数据预处理", 3)
    body(doc, "真实网络图片在尺寸、颜色通道、拍摄角度、背景复杂度和主体比例上差异较大，不能直接作为模型输入。本文将所有下载成功的图片统一转为 RGB，缩放为 64×64，按固定类别顺序编码标签，并采用分层策略划分训练集、验证集和测试集。MLP 路径使用展平像素并基于训练集均值、标准差进行标准化；CNN-local features 路径先提取固定卷积响应，再进行池化和颜色统计拼接。")
    table_doc(
        doc,
        num,
        ["步骤", "处理方法", "目的"],
        [
            ["来源筛选", "根据目标类别 MID 在官方标签表中筛选正样本", "保证类别来源明确、标签可追踪"],
            ["URL 连接", "将 ImageID 与验证集 URL 表连接", "获取缩略图、原图、作者和许可字段"],
            ["图像下载", "优先使用 Thumbnail300KURL，失败写入日志", "控制文件大小并保留失败记录"],
            ["尺寸统一", "转为 RGB 并缩放到 64×64", "形成统一模型输入"],
            ["标准化", "使用训练集均值和标准差处理特征", "降低尺度差异对训练的影响"],
            ["分层划分", "尽量保持训练/验证/测试类别比例", "减少类别分布偏差"],
        ],
        "数据预处理流程与目的",
        [1700, 3400, CONTENT_DXA - 5100],
    )

    h(doc, "（三）数据可视化", 3)
    body(doc, "数据可视化用于检查样本是否符合任务预期。样本预览图显示，Pilot 子集包含动物、交通工具和日常物体，既有轮廓清晰的样本，也有背景复杂、主体较小或视角变化明显的样本。这种真实图像复杂性正是分类难度和错误分析的来源。")
    figure(doc, num, fig_path("figure_sample_grid.png"), "Pilot 子集样本预览", "Open Images 元数据指向的公开缩略图，本项目下载后用 Python/PIL 排版")

    h(doc, "5.2 实验设置", 2)
    h(doc, "（一）实验环境", 3)
    table_doc(
        doc,
        num,
        ["项目", "环境说明"],
        [
            ["运行平台", "Windows 本地环境"],
            ["主要语言", "Python 3.x"],
            ["核心依赖", "NumPy、Pillow、matplotlib、python-docx；完整 ResidualCNN 训练需 PyTorch"],
            ["硬件条件", "课程本地电脑即可运行 NumPy 快速实验；PyTorch 扩展训练建议使用 GPU"],
            ["输出文件", "data、results、figures、logs、paper 文件夹"],
        ],
        "实验环境与运行平台",
        [2300, CONTENT_DXA - 2300],
    )

    h(doc, "（二）参数设置", 3)
    table_doc(
        doc,
        num,
        ["参数", "设置", "选择依据"],
        [
            ["图像尺寸", "64×64×3", "兼顾真实图像信息和课程环境运行速度"],
            ["类别数量", "12 类", "保证任务具有难度，同时避免小样本过度稀疏"],
            ["快速实验轮数", "80 epochs", "足以观察训练损失下降和验证波动"],
            ["MLP 隐藏层", "96 个隐藏单元", "在小数据上控制模型容量"],
            ["正则化", "L2 weight decay = 1e-4", "缓解权重过大和过拟合"],
            ["随机种子", "42", "保证划分与训练结果可重复"],
        ],
        "主要参数设置与依据",
        [2100, 2100, CONTENT_DXA - 4200],
    )

    h(doc, "（三）对照实验", 3)
    body(doc, "本文设置 MLP(flatten pixels)、CNN-local features 和 ResidualCNN 三条模型路线。前两者用于当前 Pilot 子集的快速可运行对照，第三者用于体现完整卷积网络复现入口。这样的设置可以在不牺牲课程可检查性的前提下，把基础理论、轻量实验和深层模型扩展连接起来。")
    table_doc(
        doc,
        num,
        ["模型", "输入形式", "结构要点", "实验意义"],
        [
            ["MLP(flatten pixels)", "64×64×3 展平像素", "全连接层、ReLU、Softmax", "检验不利用空间结构时的基线效果"],
            ["CNN-local features", "固定卷积响应、池化特征与颜色统计", "Sobel/Laplacian/Mean filter + MLP", "验证局部卷积特征的作用"],
            ["ResidualCNN", "RGB 图像张量", "Conv-BN-ReLU、Residual Block、GAP、Linear", "提供可扩展的完整卷积网络复现入口"],
        ],
        "三类模型的角色与差异",
        [2100, 2400, 2500, CONTENT_DXA - 7000],
    )
    figure(doc, num, fig_path("figure_model_architecture.png"), "ResidualCNN 与 MLP 的结构对照", "本项目根据模型代码用 Python/PIL 自行绘制")


def results(doc: Document, num: Numbering, metrics: dict) -> None:
    h(doc, "六、实验结果与分析", 1)
    h(doc, "（一）定量评估", 2)
    body(doc, "分类任务的定量评估采用 Accuracy、Macro-F1 和混淆矩阵。Accuracy 反映总体预测正确比例，Macro-F1 对每一类赋予相同权重，更适合观察多类别小样本任务中各类表现是否均衡。")
    model_rows = [
        [m["model"], f"{m['accuracy'] * 100:.2f}%", f"{m['macro_f1']:.4f}", "Pilot 测试集"]
        for m in metrics.get("models", [])
    ]
    table_doc(
        doc,
        num,
        ["模型", "Accuracy", "Macro-F1", "数据范围"],
        model_rows,
        "Pilot 测试集模型性能对比",
        [2800, 1500, 1500, CONTENT_DXA - 5800],
    )
    body(doc, "结果显示，MLP(flatten pixels) 的 Accuracy 为 10.00%，Macro-F1 为 0.0556；CNN-local features 的 Accuracy 为 20.00%，Macro-F1 为 0.0972。两个模型绝对数值都不高，原因是测试集只有 10 张图像且覆盖 12 类，单张图片会显著影响比例。但在相同数据划分和相同分类器容量下，卷积局部特征路径优于展平像素路径，说明空间局部特征在真实图像分类中提供了更有效的信息。")
    figure(doc, num, fig_path("figure_training_curves.png"), "训练损失与验证准确率曲线", "results/training_log.csv，本项目用 Python/PIL 自行绘制")
    figure(doc, num, fig_path("figure_metrics_comparison.png"), "模型指标对比图", "results/model_comparison.csv，本项目用 Python/PIL 自行绘制")

    h(doc, "（二）定性评估", 2)
    body(doc, "定性评估重点观察样本外观、类别混淆和错误来源。样本预览可以看出，真实图像中主体大小、背景复杂度、拍摄角度和颜色分布差异明显。固定卷积特征能够捕捉边缘与纹理，但对于遮挡、复杂背景和高层语义仍然不足。")
    body(doc, "混淆矩阵中的错误提示：部分类别容易被预测到外观或背景相近的类别上。例如，交通工具之间可能共享轮廓边缘，动物样本可能受毛发纹理和背景影响，椅子、花等物体也可能因为主体占比小而被误判。错误案例说明，当前模型更多依赖低层视觉统计，尚未形成稳定的高层语义表示。")
    figure(doc, num, fig_path("figure_confusion_matrix.png"), "CNN-local features 在 Pilot 测试集上的混淆矩阵", "results/metrics.json，本项目用 Python/PIL 自行绘制")
    table_doc(
        doc,
        num,
        ["错误类型", "可能原因", "改进方案"],
        [
            ["动物与物体混淆", "背景复杂，低层纹理不足以表达语义", "扩大样本量并训练端到端 ResidualCNN"],
            ["交通工具互相混淆", "轮廓边缘相似，视角变化较大", "加入数据增强和更深层特征"],
            ["小主体样本误判", "主体占图像比例低，背景信息干扰分类", "进行目标裁剪或使用检测辅助预处理"],
            ["少样本类别不稳定", "测试集中部分类别样本极少", "构建更均衡的训练/验证/测试集"],
        ],
        "典型错误原因与改进方向",
        [2100, 3500, CONTENT_DXA - 5600],
    )

    h(doc, "（三）代码运行截图", 2)
    body(doc, "代码运行记录用于证明程序可以从数据、训练、评估到图表生成形成完整链路。项目一键流程会依次完成 Open Images 元数据下载、Pilot 样本构建、NumPy 快速实验、指标保存、可视化生成和 Word 论文生成。运行记录图展示了主要命令、输出文件和关键指标。")
    figure(doc, num, fig_path("figure_code_run_result.png"), "项目运行过程与结果记录", "本项目运行日志、metrics.json 和输出文件整理")

    h(doc, "（四）讨论", 2)
    body(doc, "本实验的优势在于来源清楚、链路完整、对照关系明确。与直接使用玩具数据集相比，Open Images 子集更接近真实场景；与只写模型原理相比，本项目包含实际样本下载、预处理、训练日志、指标文件和可视化输出。MLP 与 CNN-local features 的比较也使卷积归纳偏置的作用更加直观。")
    body(doc, "本实验的不足同样明显。首先，Pilot 子集规模较小，测试指标方差较大；其次，CNN-local features 只能近似体现卷积思想，不能替代端到端 CNN；再次，完整 ResidualCNN 入口虽然已经提供，但若要得到更可靠的深度学习结论，需要在更大 per_class 和更强硬件环境下训练。后续工作可将每类样本扩展到 1000 张以上，加入随机裁剪、颜色扰动、水平翻转和早停策略，并与标准 ResNet18 进行对比。")


def conclusion_refs_appendix(doc: Document, num: Numbering) -> None:
    h(doc, "七、结论", 1)
    body(doc, "本文围绕真实图像分类任务，完成了从 Open Images 数据采集到卷积神经网络复现的完整课程论文项目。论文把课堂中关于 MLP、卷积、池化和残差学习的基础知识转化为一个可运行的问题链条：图像展平成向量会丢失空间结构，因此需要卷积；普通卷积网络加深后可能训练困难，因此引入残差连接；真实数据库过大不适合全部打包，因此基于官方元数据构建可追踪、可扩展的 Pilot 子集。")
    body(doc, "实验结果表明，在当前 Pilot 子集上，基于局部卷积响应的模型相对于展平像素 MLP 获得更好的 Accuracy 和 Macro-F1，说明卷积结构对图像任务具有必要性。与此同时，训练曲线和混淆矩阵也说明，小样本条件下模型容易过拟合，绝对指标不能被过度解释。因此，本文将结论限定在可由实验支持的范围内：卷积归纳偏置在真实图像分类中表现出相对优势，ResidualCNN 是进一步扩大数据和深化训练的合理方向。")
    body(doc, "最终交付项目包括爬虫代码、数据表、图像样本、快速实验代码、ResidualCNN 训练代码、实验结果、可视化图片和 Word 论文生成脚本。它的价值不仅在于得到一次分类结果，更在于把人工智能算法复现所需的关键环节完整连接起来：问题提出、数据构建、模型选择、实验验证、结果解释和局限讨论。")

    h(doc, "八、参考文献", 1)
    refs = [
        "[1] Google Research. Open Images Dataset V7 Overview. https://storage.googleapis.com/openimages/web/index.html",
        "[2] Google Research. Open Images V7 Download Page. https://storage.googleapis.com/openimages/web/download_v7.html",
        "[3] LeCun Y., Bottou L., Bengio Y., Haffner P. Gradient-based learning applied to document recognition. Proceedings of the IEEE, 1998.",
        "[4] Krizhevsky A., Sutskever I., Hinton G. ImageNet Classification with Deep Convolutional Neural Networks. NeurIPS, 2012.",
        "[5] He K., Zhang X., Ren S., Sun J. Deep Residual Learning for Image Recognition. CVPR, 2016.",
        "[6] PyTorch Contributors. PyTorch Documentation. https://pytorch.org/docs/",
        "[7] NumPy Developers. NumPy Documentation. https://numpy.org/doc/",
        "[8] Pillow Contributors. Pillow Documentation. https://pillow.readthedocs.io/",
        "[9] 课程教学资料：多层感知机、卷积神经网络、深度视觉模型与模型评估相关章节。",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        set_para(p, after=4, line=1.25, first_line=False)
        set_run_font(p.add_run(ref), 10.5, east_asia="宋体", color="000000")

    h(doc, "附录A：GitHub 源代码链接与运行指南", 1)
    body(doc, "本项目代码已经整理为可上传 GitHub 的仓库结构，包含 README、requirements.txt、src、scripts、data、results、figures、logs、paper 等目录。由于当前未提供已公开的 GitHub 仓库地址，本文不虚构链接；提交前可将代码上传至 GitHub，并把链接补充为：GitHub 源代码链接：________________。")
    table_doc(
        doc,
        num,
        ["目录或文件", "作用说明"],
        [
            ["README.md", "说明项目任务、依赖安装、运行命令、数据来源和结果指标"],
            ["requirements.txt", "记录 NumPy、Pillow、matplotlib、python-docx 等依赖"],
            ["src/openimages_crawler.py", "下载官方元数据并构建 Open Images Pilot 子集"],
            ["src/numpy_experiment.py", "运行 MLP 与 CNN-local features 快速对照实验"],
            ["src/torch_models.py / src/train_torch.py", "定义并训练完整 MLPBaseline 与 ResidualCNN"],
            ["src/visualization.py", "生成论文中的统计图、曲线图、混淆矩阵和运行记录图"],
            ["paper/", "保存生成的 Word 论文"],
        ],
        "项目结构与代码模块说明",
        [2900, CONTENT_DXA - 2900],
    )
    body(doc, "运行指南：在项目根目录执行 python -m pip install -r requirements.txt 安装依赖；执行 python main.py --all --per-class 8 可完成数据构建、快速实验、图表生成和论文生成；若只重新生成图表，可执行 python main.py --figures；若只重新生成论文，可执行 python main.py --report；完整 ResidualCNN 训练需安装 PyTorch 后执行 python -m src.train_torch。")

    h(doc, "附录B：自构数据集说明", 1)
    body(doc, "本文使用的 Pilot 子集属于基于公开数据源整理的自构数据集。数据并非直接下载一个现成压缩包，而是根据课程任务自行确定类别、编写筛选逻辑、连接官方元数据、下载图像缩略图、处理失败样本、统一尺寸和标签，并保存样本索引。该过程具有可追踪、可扩展和可复现特点。")
    table_doc(
        doc,
        num,
        ["环节", "本人完成的工作", "形成的材料"],
        [
            ["类别设计", "选择 12 个常见视觉类别并记录 MID", "类别列表与筛选参数"],
            ["数据采集", "连接 Open Images 官方 CSV 并下载缩略图", "data/raw、data/images"],
            ["清洗整理", "剔除下载失败或无法解析样本，记录失败原因", "logs/download_failures.csv"],
            ["标注索引", "保存类别、URL、作者、许可和本地路径", "openimages_pilot_images.csv"],
            ["可视化检查", "生成样本预览、数据集统计和类别分布图", "figures 文件夹"],
        ],
        "自构数据集整理过程",
        [1800, 4100, CONTENT_DXA - 5900],
    )

    h(doc, "附录C：论文格式排版说明", 1)
    body(doc, "本文按课程论文结构组织，一级标题、二级标题和三级标题层级清楚；正文采用小四号中文字体、1.5 倍行距；表格采用比正文略小的字号，英文和数字使用 Times New Roman；所有表格均设置表题，所有图片均设置图题和来源；公式置于无边框表格中并右侧编号；参考文献集中列出，文中资料来源与文末条目对应。")


def build() -> Path:
    rows = read_csv(ROOT / "data" / "processed" / "openimages_pilot_images.csv")
    metrics = read_json(ROOT / "results" / "metrics.json")
    num = Numbering()
    doc = Document()
    configure(doc)
    header_footer(doc)
    title_page(doc, num, rows, metrics)
    source_statement(doc, num)
    abstract_intro(doc, num, rows, metrics)
    algorithm(doc)
    experiment(doc, num, rows, metrics)
    results(doc, num, metrics)
    conclusion_refs_appendix(doc, num)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
